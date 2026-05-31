"""
Генератор DOCX-документов СУЖЦД.
Создаёт файлы в папке prj_docs/ на основе данных из БД.
Формат оформления — по ГОСТ (основная надпись, таблицы, разделы).
"""
from pathlib import Path
from datetime import date
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import sys, os

sys.path.insert(0, str(Path(__file__).parent))
# os.chdir не используем — меняет CWD всего процесса uvicorn,
# ломает относительные пути в других модулях.

from database import SessionLocal
from models import DocumentInstance, DocumentType

OUT = Path(__file__).parent / "prj_docs"
OUT.mkdir(exist_ok=True)


def _today() -> str:
    """Текущая дата в момент генерации документа (не при запуске сервера)."""
    return date.today().strftime("%d.%m.%Y")

# ──────────────────────────────────────────────────────────────────────────────
# Утилиты оформления
# ──────────────────────────────────────────────────────────────────────────────

def set_margins(doc, top=2, bottom=2, left=3, right=1.5):
    for section in doc.sections:
        section.top_margin    = Cm(top)
        section.bottom_margin = Cm(bottom)
        section.left_margin   = Cm(left)
        section.right_margin  = Cm(right)

def add_title(doc, text, size=14, bold=True, center=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    return p

def add_subtitle(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    return p

def add_section(doc, title, text="", size=12):
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(size)
    if text:
        pt = doc.add_paragraph(text)
        if pt.runs:
            pt.runs[0].font.size = Pt(11)
    return p

def _cell_fmt(cell, bold=False, size=10, align=None):
    """Безопасно форматирует первый параграф ячейки таблицы."""
    p = cell.paragraphs[0]
    if not p.runs:
        p.add_run(cell.text)  # принудительно создаём run если его нет
    for run in p.runs:
        if bold:
            run.bold = True
        run.font.size = Pt(size)
    if align:
        p.alignment = align


def add_field_table(doc, rows: list[tuple[str,str]]):
    """Таблица «Поле — Значение»."""
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = table.rows[0].cells
    hdr[0].text = "Параметр"
    hdr[1].text = "Значение"
    for cell in hdr:
        _cell_fmt(cell, bold=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    for label, value in rows:
        row = table.add_row().cells
        row[0].text = label
        row[1].text = str(value) if value else "—"
        for cell in row:
            _cell_fmt(cell, size=10)
    # ширина колонок
    for row in table.rows:
        row.cells[0].width = Cm(7)
        row.cells[1].width = Cm(10)
    return table

def add_stamp(doc, doc_id, designation, gost, date_str=None):
    """Основная надпись (упрощённая)."""
    doc.add_paragraph()
    table = doc.add_table(rows=3, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    date_str = date_str or _today()
    data = [
        ["Обозначение", designation, "ГОСТ", gost],
        ["ID документа", doc_id,     "Дата", date_str],
        ["Разработал",   "СУЖЦД v1.0","Лист", "1"],
    ]
    for i, row_data in enumerate(data):
        cells = table.rows[i].cells
        for j, val in enumerate(row_data):
            cells[j].text = val
            for p in cells[j].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(8)
                    run.bold = (j % 2 == 0)

def finalize(doc, path: Path, verbose: bool = False):
    doc.save(path)
    if verbose:
        print(f"  ✓ {path.name}")


# ──────────────────────────────────────────────────────────────────────────────
# Генераторы по типу документа
# ──────────────────────────────────────────────────────────────────────────────

def gen_eskd_detail(d, fields, path):
    doc = Document()
    set_margins(doc)
    add_title(doc, "ЧЕРТЁЖ ДЕТАЛИ", 15)
    add_subtitle(doc, f"ГОСТ 2.102–2013 · ЕСКД")
    doc.add_paragraph()
    add_title(doc, fields.get("name_field", d.name), 13)
    add_subtitle(doc, fields.get("designation", d.designation))
    doc.add_paragraph()
    add_section(doc, "1. Основные параметры детали")
    add_field_table(doc, [
        ("Обозначение",              fields.get("designation",    d.designation)),
        ("Наименование детали",      fields.get("name_field",     d.name)),
        ("Материал",                 fields.get("material",       "—")),
        ("Масса, кг",                fields.get("mass",           "—")),
        ("Масштаб",                  fields.get("scale",          "—")),
        ("Допуски формы и расположения", fields.get("tolerance",  "—")),
        ("Шероховатость",            fields.get("roughness",      "—")),
        ("Покрытие",                 fields.get("coating",        "—")),
        ("Термообработка",           fields.get("heat_treatment", "—")),
    ])
    doc.add_paragraph()
    if fields.get("tech_requirements"):
        add_section(doc, "2. Технические требования", fields["tech_requirements"])
    else:
        add_section(doc, "2. Технические требования")
        doc.add_paragraph(
            f"1. Точность изготовления по ГОСТ 25346–89.\n"
            f"2. Материал {fields.get('material','—')} ГОСТ 380–2005.\n"
            f"3. Покрытие: {fields.get('coating','—')}.\n"
            f"4. {fields.get('heat_treatment','Без термообработки')}.\n"
            f"5. Неуказанные предельные отклонения: H14, h14, ±IT14/2."
        ).runs[0].font.size = Pt(11)
    add_section(doc, "3. Каскадные зависимости (ГОСТ Р 2.503–2023)")
    doc.add_paragraph(
        "• Изменение «material» → ESTD_MK (material, material_rate) [обязательная]\n"
        "• Изменение «tolerance» → ESTD_OK (cutting_modes, tool) [обязательная]\n"
        "• Изменение «roughness» → ESTD_OK (finishing_method) [обязательная]\n"
        "• Изменение «mass»      → ESKD_PASSPORT (tech_chars) [обязательная]"
    ).runs[0].font.size = Pt(10)
    add_stamp(doc, d.id, d.designation, "ГОСТ 2.102–2013")
    finalize(doc, path)


def gen_eskd_spec(d, fields, path):
    doc = Document()
    set_margins(doc)
    add_title(doc, "СПЕЦИФИКАЦИЯ", 15)
    add_subtitle(doc, "ГОСТ 2.102–2013 · ЕСКД")
    doc.add_paragraph()
    add_title(doc, d.name, 13)
    add_subtitle(doc, d.designation)
    doc.add_paragraph()
    add_section(doc, "Состав изделия")
    # Парсим состав
    comp = fields.get("composition", "")
    if comp:
        p = doc.add_paragraph()
        p.runs  # ensure
        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text = "№ поз."
        hdr[1].text = "Наименование"
        hdr[2].text = "Кол-во"
        for h in hdr:
            _cell_fmt(h, bold=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
        # Разбиваем состав по запятой/точке с запятой
        parts = [p.strip() for p in comp.replace(";",",").split(",") if p.strip()]
        for i, part in enumerate(parts, 1):
            row = table.add_row().cells
            row[0].text = str(i)
            row[1].text = part
            row[2].text = "1"
            for c in row:
                _cell_fmt(c, size=10)
    doc.add_paragraph()
    add_field_table(doc, [
        ("Обозначение спецификации", d.designation),
        ("Количество изделий",       fields.get("quantity", "1")),
    ])
    add_stamp(doc, d.id, d.designation, "ГОСТ 2.102–2013")
    finalize(doc, path)


def gen_estd_mk(d, fields, path):
    doc = Document()
    set_margins(doc)
    add_title(doc, "МАРШРУТНАЯ КАРТА", 15)
    add_subtitle(doc, "ГОСТ 3.1118–82 · ЕСТД")
    doc.add_paragraph()
    add_title(doc, fields.get("product_name", d.name), 13)
    doc.add_paragraph()
    add_section(doc, "1. Общие сведения об изделии")
    add_field_table(doc, [
        ("Наименование изделия",        fields.get("product_name",  d.name)),
        ("Обозначение КД",              fields.get("kd_designation","—")),
        ("Материал (строка М)",         fields.get("material",      "—")),
        ("Масса заготовки, кг",         fields.get("blank_mass",    "—")),
        ("Норма расхода материала, кг", fields.get("material_rate", "—")),
        ("Оборудование",                fields.get("equipment",     "—")),
    ])
    doc.add_paragraph()
    add_section(doc, "2. Маршрут обработки (операции)")
    ops = fields.get("operations", "")
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Код опер."
    hdr[1].text = "Наименование операции"
    hdr[2].text = "Оборудование"
    for h in hdr:
        _cell_fmt(h, bold=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    for op in [o.strip() for o in ops.split(",") if o.strip()]:
        parts = op.split(" ", 1)
        row = table.add_row().cells
        row[0].text = parts[0] if len(parts) > 1 else "—"
        row[1].text = parts[1] if len(parts) > 1 else parts[0]
        row[2].text = fields.get("equipment", "—")
        for c in row:
            _cell_fmt(c, size=10)
    add_stamp(doc, d.id, d.designation, "ГОСТ 3.1118–82")
    finalize(doc, path)


def gen_estd_ok(d, fields, path):
    doc = Document()
    set_margins(doc)
    add_title(doc, "ОПЕРАЦИОННАЯ КАРТА", 15)
    add_subtitle(doc, "ГОСТ 3.1404–86 · ЕСТД")
    doc.add_paragraph()
    add_title(doc, f"Операция {fields.get('op_code','—')} — {fields.get('op_name','—')}", 13)
    doc.add_paragraph()
    add_section(doc, "1. Параметры операции")
    add_field_table(doc, [
        ("Код операции",            fields.get("op_code",          "—")),
        ("Наименование операции",   fields.get("op_name",          "—")),
        ("Оборудование",            fields.get("equipment",        "—")),
        ("Режимы резания",          fields.get("cutting_modes",    "—")),
        ("Инструмент",              fields.get("tool",             "—")),
        ("Приспособление",          fields.get("fixture",          "—")),
        ("Метод финишной обработки",fields.get("finishing_method", "—")),
        ("Режимы термообработки",   fields.get("heat_modes",       "—")),
        ("Норма времени, мин",      fields.get("norm_time",        "—")),
    ])
    doc.add_paragraph()
    add_section(doc, "2. Переходы операции")
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for h, t in zip(hdr, ["№", "Содержание перехода", "Инструмент"]):
        h.text = t
        _cell_fmt(h, bold=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    transitions = [
        ("1", f"Установить и закрепить деталь. {fields.get('fixture','')}",
         "—"),
        ("2", f"Обработать поверхность. {fields.get('cutting_modes','')}",
         fields.get("tool", "—")),
        ("3", f"Финишная обработка. {fields.get('finishing_method','')}",
         fields.get("tool", "—")),
    ]
    for num, content, tool in transitions:
        row = table.add_row().cells
        row[0].text = num
        row[1].text = content
        row[2].text = tool
        for c in row:
            _cell_fmt(c, size=10)
    add_stamp(doc, d.id, d.designation, "ГОСТ 3.1404–86")
    finalize(doc, path)


def gen_espd_tz(d, fields, path):
    doc = Document()
    set_margins(doc)
    add_title(doc, "ТЕХНИЧЕСКОЕ ЗАДАНИЕ", 15)
    add_subtitle(doc, "ГОСТ 19.201–78 · ЕСПД")
    doc.add_paragraph()
    add_title(doc, d.name, 13)
    add_subtitle(doc, d.designation)
    doc.add_paragraph()

    sections = [
        ("1. Назначение и область применения",
         fields.get("purpose", "—")),
        ("2. Функциональные требования",
         fields.get("func_requirements", "—")),
        ("3. Требования к надёжности",
         fields.get("reliability_req", "—")),
        ("4. Требования к техническим средствам",
         fields.get("hw_req", "—")),
        ("5. Требования к программному обеспечению",
         fields.get("sw_req", "Не установлены.")),
        ("6. Требования к документации",
         fields.get("doc_req",
                    "Техническое задание, описание программы, руководство оператора.")),
        ("7. Стадии и этапы разработки",
         fields.get("stages", "—")),
    ]
    for title, text in sections:
        add_section(doc, title, text)
        doc.add_paragraph()

    add_section(doc, "8. Каскадные зависимости (ГОСТ Р 2.503–2023)")
    doc.add_paragraph(
        "• func_requirements → ESPD_OP.functions     [обязательная, Ω₂/Ω₃]\n"
        "• func_requirements → ESPD_RO.description   [обязательная, Ω₂/Ω₃]\n"
        "• func_requirements → ESPD_PMI.test_methods [обязательная, Ω₂]\n"
        "• reliability_req   → ESPD_PMI.test_methods [обязательная, Ω₂/Ω₄]\n"
        "• stages            → ESPD_OP.general_info  [рекомендуемая, Ω₃]"
    ).runs[0].font.size = Pt(10)
    add_stamp(doc, d.id, d.designation, "ГОСТ 19.201–78")
    finalize(doc, path)


def gen_eskd_passport(d, fields, path):
    doc = Document()
    set_margins(doc)
    add_title(doc, "ПАСПОРТ", 15)
    add_subtitle(doc, "ГОСТ 2.102–2013 · ЕСКД")
    doc.add_paragraph()
    add_title(doc, d.name, 13)
    add_subtitle(doc, d.designation)
    doc.add_paragraph()
    add_section(doc, "1. Технические характеристики")
    tc = fields.get("tech_chars", "—")
    for line in tc.split(","):
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(line.strip()).font.size = Pt(11)
    doc.add_paragraph()
    add_section(doc, "2. Комплект поставки")
    ds = fields.get("delivery_set", "—")
    for item in ds.split(","):
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item.strip()).font.size = Pt(11)
    doc.add_paragraph()
    add_section(doc, "3. Гарантийные обязательства",
                fields.get("guarantee", "—"))
    doc.add_paragraph()
    add_section(doc, "4. Свидетельство о приёмке")
    add_field_table(doc, [
        ("Обозначение изделия", d.designation),
        ("Дата выпуска",        _today()),
        ("Заключение ОТК",      "Соответствует ТУ. Допущен к эксплуатации."),
    ])
    add_stamp(doc, d.id, d.designation, "ГОСТ 2.102–2013")
    finalize(doc, path)


def gen_eskd_tu(d, fields, path):
    doc = Document()
    set_margins(doc)
    add_title(doc, "ТЕХНИЧЕСКИЕ УСЛОВИЯ", 15)
    add_subtitle(doc, "ГОСТ 2.114–2016 · ЕСКД")
    doc.add_paragraph()
    add_title(doc, d.name, 13)
    add_subtitle(doc, d.designation)
    doc.add_paragraph()
    sections = [
        ("1. Технические требования",  fields.get("tech_req",        "—")),
        ("2. Правила приёмки",         fields.get("acceptance",      "—")),
        ("3. Методы контроля",         fields.get("control_methods", "—")),
        ("4. Гарантийный срок",        fields.get("guarantee",       "—")),
        ("5. Условия эксплуатации",    fields.get("conditions",      "—")),
    ]
    for title, text in sections:
        add_section(doc, title, text)
        doc.add_paragraph()
    add_section(doc, "6. Каскадные зависимости (ГОСТ Р 2.503–2023)")
    doc.add_paragraph(
        "• tech_req      → ESTD_KTK.control_methods [обязательная, Ω₂/Ω₄]\n"
        "• guarantee     → ESKD_PASSPORT.guarantee  [обязательная, Ω₄]\n"
        "• conditions    → ESKD_RE.conditions        [обязательная, Ω₂/Ω₄]\n"
        "• control_methods → ESTD_KTK.control_methods [обязательная, Ω₂]"
    ).runs[0].font.size = Pt(10)
    add_stamp(doc, d.id, d.designation, "ГОСТ 2.114–2016")
    finalize(doc, path)


# ──────────────────────────────────────────────────────────────────────────────
# Диспетчер
# ──────────────────────────────────────────────────────────────────────────────

GENERATORS = {
    "ESKD_DETAIL":   gen_eskd_detail,
    "ESKD_SPEC":     gen_eskd_spec,
    "ESTD_MK":       gen_estd_mk,
    "ESTD_OK":       gen_estd_ok,
    "ESPD_TZ":       gen_espd_tz,
    "ESKD_PASSPORT": gen_eskd_passport,
    "ESKD_TU":       gen_eskd_tu,
}

def safe_filename(name: str) -> str:
    return "".join(c if c.isalnum() or c in "._- " else "_" for c in name).strip()


def generate_docx(doc: "DocumentInstance") -> bytes:
    """
    Генерирует DOCX-файл для документа и возвращает содержимое в виде байт.
    Используется для сохранения в хранилище версий (doc_storage).
    Если тип документа не поддерживается — создаёт универсальный шаблон.
    """
    import io
    import tempfile

    gen = GENERATORS.get(doc.doc_type)
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp_path = Path(tmp.name)

    try:
        if gen:
            gen(doc, doc.fields_json or {}, tmp_path)
        else:
            _gen_generic(doc, tmp_path)

        with open(tmp_path, "rb") as f:
            return f.read()
    finally:
        tmp_path.unlink(missing_ok=True)


def _gen_generic(doc: "DocumentInstance", path: Path):
    """Универсальный генератор для типов без специализированного шаблона."""
    d = Document()
    d.add_heading(doc.name, level=1)
    d.add_paragraph(f"Тип документа: {doc.doc_type}")
    d.add_paragraph(f"Обозначение: {doc.designation}")
    d.add_paragraph(f"Версия: {doc.version}")
    d.add_paragraph(f"Статус: {doc.status}")
    d.add_paragraph(f"Ветка: {doc.branch_name}")
    d.add_heading("Поля документа", level=2)
    fields = doc.fields_json or {}
    if fields:
        tbl = d.add_table(rows=1, cols=2)
        tbl.style = "Table Grid"
        hdr = tbl.rows[0].cells
        hdr[0].text = "Поле"
        hdr[1].text = "Значение"
        for fid, val in fields.items():
            row = tbl.add_row().cells
            row[0].text = str(fid)
            row[1].text = str(val)
    else:
        d.add_paragraph("Поля не заполнены.")
    add_stamp(d, doc.id, doc.designation, "СУЖЦД")
    finalize(d, path)


def main():
    db = SessionLocal()
    docs = db.query(DocumentInstance).all()
    print(f"\nГенерирую документы в {OUT}/\n")
    generated = 0
    skipped = 0
    for d in docs:
        fields = d.fields_json or {}
        gen = GENERATORS.get(d.doc_type)
        if not gen:
            print(f"  ⊘ {d.id} ({d.doc_type}) — генератор не реализован, пропуск")
            skipped += 1
            continue
        # Пропускаем тестовый документ без полноценных данных
        if d.id == "DOC-TZ-001":
            print(f"  ⊘ {d.id} — тестовый документ, пропуск")
            skipped += 1
            continue
        fname = f"{d.id}_{safe_filename(d.name)}.docx"
        path = OUT / fname
        try:
            gen(d, fields, path)
            print(f"  ✓ {fname}")
            generated += 1
        except Exception as e:
            print(f"  ✗ {d.id}: {e}")
            skipped += 1
    db.close()
    print(f"\nИтого: создано {generated}, пропущено {skipped}")
    print(f"Папка: {OUT.resolve()}")

if __name__ == "__main__":
    main()
