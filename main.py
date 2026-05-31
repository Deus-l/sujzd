"""
СУЖЦД — Система Управления Жизненным Циклом Документации
FastAPI-приложение: все REST-маршруты + статические файлы
"""
import json
import os
from datetime import datetime
from pathlib import Path
from typing import List, Optional

from fastapi import FastAPI, Depends, HTTPException, Query, UploadFile, File, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import HTMLResponse, JSONResponse, FileResponse, Response
from fastapi.staticfiles import StaticFiles
from sqlalchemy.orm import Session

import models
import schemas
from database import engine, get_db, init_db
import doc_storage
from delta_repo import (
    commit_delta, checkout, diff, blame,
    create_branch, merge_branch, snapshot,
)
from dependency_engine import (
    process_delta, get_full_cascade_tree,
    get_cascade, get_matrix_stats, reload_matrix, _matrix,
)
from ia_module import (
    classify_omega, estimate_significance,
    analyze_delta_ai, parse_step_delta, parse_xlsx_delta, parse_docx_delta,
    detect_doc_type, extract_fields_from_text,
)

# ---------------------------------------------------------------------------
# Инициализация приложения
# ---------------------------------------------------------------------------

app = FastAPI(
    title="СУЖЦД — Система управления жизненным циклом документации",
    description=(
        "Реализация концепции атомарного версионирования технической документации "
        "ЕСКД/ЕСТД/ЕСПД на принципах Git с каскадным распространением изменений. "
        "ГОСТ Р 2.503–2023, ГОСТ 2.102–2013, ГОСТ 3.1102–2011, ГОСТ 19.101–77."
    ),
    version="1.0.0",
    docs_url="/api/docs",
    redoc_url="/api/redoc",
)

_CORS_ORIGINS = [
    "http://localhost:8877",
    "http://127.0.0.1:8877",
    "http://0.0.0.0:8877",
]
app.add_middleware(
    CORSMiddleware,
    allow_origins=_CORS_ORIGINS,
    allow_methods=["GET", "POST", "PUT", "PATCH", "DELETE", "OPTIONS"],
    allow_headers=["*"],
    allow_credentials=False,
)


# ---------------------------------------------------------------------------
# Инициализация БД и загрузка начальных данных
# ---------------------------------------------------------------------------

@app.on_event("startup")
def startup():
    init_db()
    db = next(get_db())
    try:
        _seed_db(db)
    except Exception as exc:
        print(f"[startup] Ошибка при инициализации БД: {exc}")
        db.rollback()
    finally:
        db.close()
    # Сгенерировать DOCX для всех документов у которых ещё нет файла в хранилище
    _ensure_all_docs_have_files()


def _seed_db(db: Session):
    # Загрузка DSR
    dsr_path = Path(__file__).parent / "data" / "dsr.json"
    with open(dsr_path, encoding="utf-8") as f:
        dsr = json.load(f)

    for code, info in dsr.items():
        if not db.query(models.DocumentType).filter_by(code=code).first():
            db.add(models.DocumentType(
                code=code,
                std=info["std"],
                short_code=info["short_code"],
                name=info["name"],
                gost=info.get("gost", ""),
                fields_json=info.get("fields", []),
            ))

    # Начальные ветки
    if not db.query(models.Branch).filter_by(name="main").first():
        db.add(models.Branch(name="main", description="Основная ветка"))

    # Демонстрационные документы
    demo_docs = [
        {"id": "DOC-001", "doc_type": "ESKD_DETAIL", "name": "Корпус датчика Д-47М",
         "designation": "СУГИ.301847.023", "version": "2.0",
         "fields_json": {"designation": "СУГИ.301847.023", "name_field": "Корпус датчика Д-47М",
                         "material": "Ст3", "mass": "0.312", "scale": "1:1",
                         "tolerance": "H8/g6", "roughness": "Ra 3.2",
                         "coating": "Цинкование Хц.6", "heat_treatment": "Нормализация"}},
        {"id": "DOC-002", "doc_type": "ESKD_SPEC", "name": "Спецификация блока управления БУ-12",
         "designation": "СУГИ.467823.012СП", "version": "1.2",
         "fields_json": {"designation": "СУГИ.467823.012СП",
                         "composition": "11 позиций (транзисторы КТ315Б ×4, резисторы R1–R6, конденсаторы C1–C3)",
                         "quantity": "1"}},
        {"id": "DOC-003", "doc_type": "ESTD_MK", "name": "МК корпуса Д-47М",
         "designation": "МК-001", "version": "1.8",
         "fields_json": {"product_name": "Корпус датчика Д-47М", "kd_designation": "СУГИ.301847.023",
                         "material": "Ст3", "blank_mass": "0.45", "material_rate": "0.40",
                         "operations": "005 Токарная, 010 Фрезерная, 020 Шлифовальная",
                         "equipment": "Токарный 16К20"}},
        {"id": "DOC-004", "doc_type": "ESTD_OK", "name": "ОК токарная операция 005",
         "designation": "ОК-005", "version": "1.5",
         "fields_json": {"op_code": "005", "op_name": "Токарная",
                         "equipment": "16К20", "cutting_modes": "n=630 об/мин, S=0.25 мм/об",
                         "tool": "Резец Т5К10 ГОСТ 18878-73", "fixture": "Патрон трёхкулачковый",
                         "finishing_method": "Шлифование Ra 3.2", "heat_modes": "Нормализация 850°С",
                         "norm_time": "12.5"}},
        {"id": "DOC-005", "doc_type": "ESPD_TZ", "name": "ТЗ на систему мониторинга СМ-4",
         "designation": "СМ-4.ТЗ.001", "version": "1.0",
         "fields_json": {"purpose": "Мониторинг параметров технологического оборудования",
                         "func_requirements": "4 аналоговых канала, частота дискретизации 500 Гц, интерфейс RS-485",
                         "reliability_req": "MTBF ≥ 3000 ч, вероятность безотказной работы 0.95",
                         "hw_req": "Процессор Cortex-M4, RAM 256 кБ, Flash 1 МБ",
                         "stages": "Аванпроект, технический проект, рабочая документация"}},
        {"id": "DOC-006", "doc_type": "ESKD_PASSPORT", "name": "Паспорт блока БУ-12",
         "designation": "СУГИ.467823.012ПС", "version": "1.1",
         "fields_json": {"tech_chars": "Uпит: 220В±10%, Pпотр: 15Вт, Масса: 1.2кг, Uвых: 5В/12В",
                         "delivery_set": "Блок БУ-12 — 1 шт., РЭ — 1 экз., формуляр — 1 экз.",
                         "guarantee": "18 месяцев со дня ввода в эксплуатацию"}},
        {"id": "DOC-007", "doc_type": "ESKD_TU", "name": "ТУ на корпус датчика Д-47М",
         "designation": "СУГИ.301847.023ТУ", "version": "1.0",
         "fields_json": {"tech_req": "Точность обработки по 7-му квалитету. Материал Ст3 ГОСТ 380-2005.",
                         "acceptance": "100% визуальный контроль, выборочный контроль размеров 10%",
                         "control_methods": "Визуальный, измерительный (штангенциркуль, нутромер)",
                         "guarantee": "12 месяцев",
                         "conditions": "Температура –40…+60°С, влажность до 95%"}},
    ]

    for d in demo_docs:
        if not db.query(models.DocumentInstance).filter_by(id=d["id"]).first():
            db.add(models.DocumentInstance(**d))

    db.commit()


def _ensure_all_docs_have_files():
    """
    При старте проверяет каждый документ в БД.
    Если для него нет DOCX-файла в хранилище — генерирует и сохраняет.
    Гарантирует что prj_docs/{doc_id}/v{version}.docx всегда существует.
    """
    from generate_docs import generate_docx
    db = next(get_db())
    try:
        docs = db.query(models.DocumentInstance).all()
        for doc in docs:
            # Проверяем: есть ли уже файл HEAD-версии?
            existing = doc_storage.get_version_path(doc.id, doc.version)
            if existing is not None:
                continue  # файл уже есть
            # Нет файла — генерируем
            try:
                file_bytes = generate_docx(doc)
                doc_storage.save_document_version(
                    doc_id=doc.id,
                    doc_name=doc.name,
                    doc_type=doc.doc_type,
                    version=doc.version,
                    docx_bytes=file_bytes,
                    author="system",
                    reason="Начальная генерация при запуске",
                    committed_at=doc.created_at,
                )
            except Exception as e:
                print(f"[doc_storage] Ошибка генерации для {doc.id}: {e}")
    except Exception as exc:
        print(f"[startup] Ошибка в _ensure_all_docs_have_files: {exc}")
    finally:
        db.close()


# ---------------------------------------------------------------------------
# Статические файлы
# ---------------------------------------------------------------------------

static_path = Path(__file__).parent / "static"
static_path.mkdir(exist_ok=True)
app.mount("/static", StaticFiles(directory=str(static_path)), name="static")


@app.get("/", response_class=HTMLResponse, include_in_schema=False)
def root():
    index = static_path / "index.html"
    return HTMLResponse(index.read_text(encoding="utf-8"))


# ===========================================================================
# API: Document Schema Registry (DSR)
# ===========================================================================

@app.get("/api/dsr", tags=["DSR — Реестр схем"])
def get_dsr(db: Session = Depends(get_db)):
    """Полный реестр типов документов ЕСКД/ЕСТД/ЕСПД"""
    types = db.query(models.DocumentType).all()
    return [
        {"code": t.code, "std": t.std, "short_code": t.short_code,
         "name": t.name, "gost": t.gost, "fields": t.fields_json}
        for t in types
    ]


@app.get("/api/dsr/{code}", tags=["DSR — Реестр схем"])
def get_dsr_type(code: str, db: Session = Depends(get_db)):
    """Схема конкретного типа документа"""
    dt = db.query(models.DocumentType).filter_by(code=code).first()
    if not dt:
        raise HTTPException(404, f"Тип документа '{code}' не найден в DSR")
    return {"code": dt.code, "std": dt.std, "short_code": dt.short_code,
            "name": dt.name, "gost": dt.gost, "fields": dt.fields_json}


@app.get("/api/dsr/{code}/fields", tags=["DSR — Реестр схем"])
def get_dsr_fields(code: str, db: Session = Depends(get_db)):
    """Поля конкретного типа документа из DSR (раздел 4.3)"""
    dt = db.query(models.DocumentType).filter_by(code=code).first()
    if not dt:
        raise HTTPException(404, f"Тип '{code}' не найден в DSR")
    fields = dt.fields_json or []
    # Обогащаем: добавляем кол-во правил матрицы по каждому полю
    for f in fields:
        f["matrix_rules_count"] = sum(
            1 for r in _matrix
            if r["source_type"] == code and r["source_field"] == f["id"]
        )
    return {"code": code, "name": dt.name, "gost": dt.gost, "fields": fields}


@app.post("/api/dsr", tags=["DSR — Реестр схем"], status_code=201)
def create_dsr_type(data: dict, db: Session = Depends(get_db)):
    """
    Добавить новый тип документа в DSR (раздел 4.3).
    Обязательные поля: code, std, short_code, name.
    """
    required = ["code", "std", "short_code", "name"]
    for r in required:
        if not data.get(r):
            raise HTTPException(400, f"Поле '{r}' обязательно")
    if data["std"] not in ("ЕСКД", "ЕСТД", "ЕСПД"):
        raise HTTPException(400, "std должен быть ЕСКД | ЕСТД | ЕСПД")
    # Валидация типа: fields должен быть списком (dict'ов), не строкой
    fields_val = data.get("fields", [])
    if not isinstance(fields_val, list):
        raise HTTPException(400, "Поле 'fields' должно быть списком (array)")
    for i, fld in enumerate(fields_val):
        if not isinstance(fld, dict) or "id" not in fld or "name" not in fld:
            raise HTTPException(400, f"Элемент fields[{i}] должен быть объектом с полями 'id' и 'name'")
    if db.query(models.DocumentType).filter_by(code=data["code"]).first():
        raise HTTPException(409, f"Тип '{data['code']}' уже существует")
    dt = models.DocumentType(
        code=data["code"],
        std=data["std"],
        short_code=data["short_code"],
        name=data["name"],
        gost=data.get("gost", ""),
        fields_json=data.get("fields", []),
    )
    db.add(dt)
    db.commit()
    db.refresh(dt)
    return {"code": dt.code, "std": dt.std, "short_code": dt.short_code,
            "name": dt.name, "gost": dt.gost, "fields": dt.fields_json}


# ===========================================================================
# API: Документы
# ===========================================================================

@app.get("/api/documents", response_model=List[schemas.DocumentOut], tags=["Документы"])
def list_documents(
    std:    Optional[str] = Query(None, description="Фильтр по стандарту: ЕСКД | ЕСТД | ЕСПД"),
    status: Optional[str] = Query(None),
    branch: Optional[str] = Query(None),
    limit:  int           = Query(100, le=1000, description="Кол-во записей"),
    offset: int           = Query(0,            description="Смещение для пагинации"),
    db: Session = Depends(get_db),
):
    q = db.query(models.DocumentInstance)
    if std:
        type_codes = [t.code for t in db.query(models.DocumentType).filter_by(std=std).all()]
        q = q.filter(models.DocumentInstance.doc_type.in_(type_codes))
    if status:
        q = q.filter_by(status=status)
    if branch:
        q = q.filter_by(branch_name=branch)
    return q.offset(offset).limit(limit).all()


@app.post("/api/documents", response_model=schemas.DocumentOut, status_code=201, tags=["Документы"])
def create_document(data: schemas.DocumentCreate, db: Session = Depends(get_db)):
    if db.query(models.DocumentInstance).filter_by(id=data.id).first():
        raise HTTPException(409, f"Документ '{data.id}' уже существует")
    dt = db.query(models.DocumentType).filter_by(code=data.doc_type).first()
    if not dt:
        raise HTTPException(404, f"Тип документа '{data.doc_type}' не найден в DSR")
    doc = models.DocumentInstance(**data.model_dump())
    db.add(doc); db.commit(); db.refresh(doc)
    # Сгенерировать и сохранить начальный DOCX в хранилище версий
    _generate_and_store(doc, reason="Создание документа")
    return doc


@app.post("/api/documents/upload", tags=["Документы"])
async def upload_document(
    file: UploadFile = File(...),
    doc_id: str = Form(...),
    doc_type: str = Form(...),
    branch_name: str = Form(default="main"),
    db: Session = Depends(get_db),
):
    """
    Загрузить PDF, DOCX или XLSX и создать запись документа.
    Из файла автоматически извлекаются все DSR-поля: материал, допуски,
    размеры, требования и пр. После создания полностью работают дельты и каскад.
    Поддерживаемые форматы: .docx, .pdf, .xlsx
    """
    import io

    ext = Path(file.filename).suffix.lower()
    if ext not in (".pdf", ".docx", ".xlsx"):
        raise HTTPException(400, "Поддерживаются только файлы .pdf, .docx и .xlsx")

    if db.query(models.DocumentInstance).filter_by(id=doc_id).first():
        raise HTTPException(409, f"Документ '{doc_id}' уже существует")

    dt = db.query(models.DocumentType).filter_by(code=doc_type).first()
    if not dt:
        raise HTTPException(404, f"Тип '{doc_type}' не найден в DSR")

    content = await file.read()
    raw_text = ""
    name = Path(file.filename).stem
    designation = doc_id
    meta_info: dict = {"source_file": file.filename, "file_format": ext.lstrip(".")}
    xlsx_rows: list = []   # структурированные строки из XLSX (для field extraction)

    # ── Извлечение текста ─────────────────────────────────────────────
    if ext == ".docx":
        try:
            from docx import Document as DocxDocument
            doc_obj = DocxDocument(io.BytesIO(content))
            props = doc_obj.core_properties
            if props.title:
                name = props.title
            if props.subject:
                designation = props.subject
            if props.author:
                meta_info["author"] = props.author

            # 1) Параграфы
            paras = [p.text.strip() for p in doc_obj.paragraphs if p.text.strip()]

            # 2) Таблицы (ключевое дополнение)
            table_lines = []
            for tbl_idx, table in enumerate(doc_obj.tables):
                headers = []
                for row_idx, row in enumerate(table.rows):
                    cells = [c.text.strip() for c in row.cells]
                    # Удаляем дубли соседних ячеек (merged cells)
                    deduped = []
                    for c in cells:
                        if not deduped or c != deduped[-1]:
                            deduped.append(c)
                    cells = deduped

                    if row_idx == 0:
                        # Первая строка — возможные заголовки
                        headers = cells
                        line = " | ".join(c for c in cells if c)
                    else:
                        if headers and len(headers) == len(cells):
                            # Формируем "Заголовок: Значение" для лучшего regex-матчинга
                            pairs = [f"{h}: {v}" for h, v in zip(headers, cells) if h and v]
                            line = "  ".join(pairs)
                            # Сохраняем структурированную строку для XLSX-подобного извлечения
                            xlsx_rows.append(dict(zip(headers, cells)))
                        else:
                            line = " | ".join(c for c in cells if c)
                    if line.strip():
                        table_lines.append(line)

            raw_text = "\n".join(paras)
            if table_lines:
                raw_text += "\n\n[ТАБЛИЦЫ]\n" + "\n".join(table_lines)

            if not props.title and paras:
                name = paras[0][:120]
            meta_info["pages"] = len(doc_obj.sections)
            meta_info["paragraphs"] = len(paras)
            meta_info["tables"] = len(doc_obj.tables)
            meta_info["table_rows"] = len(xlsx_rows)
        except Exception as e:
            meta_info["parse_warning"] = str(e)

    elif ext == ".pdf":
        try:
            from pypdf import PdfReader
            reader = PdfReader(io.BytesIO(content))
            pdf_meta = reader.metadata or {}
            if pdf_meta.get("/Title"):
                name = pdf_meta["/Title"]
            if pdf_meta.get("/Subject"):
                designation = pdf_meta["/Subject"]
            if pdf_meta.get("/Author"):
                meta_info["author"] = pdf_meta["/Author"]
            pages_text = []
            for page in reader.pages:
                t = page.extract_text()
                if t:
                    pages_text.append(t.strip())
            raw_text = "\n".join(pages_text)
            meta_info["pages"] = len(reader.pages)
        except Exception as e:
            meta_info["parse_warning"] = str(e)

    elif ext == ".xlsx":
        try:
            import openpyxl
            wb = openpyxl.load_workbook(io.BytesIO(content), data_only=True)
            text_parts = []
            all_rows: list = []
            for sheet in wb.worksheets:
                meta_info.setdefault("sheets", []).append(sheet.title)
                rows = list(sheet.iter_rows(values_only=True))
                if not rows:
                    continue
                # Первая непустая строка → заголовки
                headers = []
                data_start = 0
                for ri, row in enumerate(rows):
                    non_empty = [str(c).strip() for c in row if c is not None and str(c).strip()]
                    if non_empty:
                        headers = [str(c).strip() if c is not None else "" for c in row]
                        data_start = ri + 1
                        break
                if not headers:
                    continue

                sheet_rows = []
                for row in rows[data_start:]:
                    cells = [str(c).strip() if c is not None else "" for c in row]
                    if any(cells):  # пропускаем пустые строки
                        row_dict = {h: v for h, v in zip(headers, cells) if h}
                        sheet_rows.append(row_dict)
                        # Строка для text-извлечения
                        pairs = [f"{h}: {v}" for h, v in row_dict.items() if v]
                        if pairs:
                            text_parts.append("  ".join(pairs))
                all_rows.extend(sheet_rows)

            xlsx_rows = all_rows
            raw_text = "\n".join(text_parts)

            # Попытка взять название из первой значимой ячейки A1
            if wb.worksheets:
                a1 = wb.worksheets[0].cell(1, 1).value
                if a1:
                    name = str(a1)[:120]
            meta_info["sheets_count"] = len(wb.worksheets)
            meta_info["total_data_rows"] = len(all_rows)
        except Exception as e:
            meta_info["parse_warning"] = str(e)

    # ── Автоопределение типа ───────────────────────────────────────────
    auto_detected = detect_doc_type(raw_text) if raw_text else None
    if auto_detected and auto_detected != doc_type:
        meta_info["_auto_detected_type"] = auto_detected

    # ── Умное извлечение DSR-полей ────────────────────────────────────
    fields_json: dict = {}
    if raw_text:
        fields_json = extract_fields_from_text(raw_text, doc_type)

    # Дополнительное извлечение по структурированным строкам (таблицы DOCX / XLSX)
    if xlsx_rows:
        from ia_module import extract_fields_from_rows
        table_fields = extract_fields_from_rows(xlsx_rows, doc_type)
        # Табличные данные приоритетнее текстовых (точнее)
        fields_json.update(table_fields)

    # Перезаписываем designation/name если нашли в тексте
    if "designation" in fields_json and len(fields_json.get("designation", "")) > 3:
        designation = fields_json["designation"]
    if "name_field" in fields_json and len(fields_json.get("name_field", "")) > 3:
        name = fields_json["name_field"]

    # Добавляем мета-информацию (не перебиваем извлечённые поля)
    fields_json.update({k: v for k, v in meta_info.items() if k not in fields_json})
    fields_json["designation"] = designation

    doc = models.DocumentInstance(
        id=doc_id,
        doc_type=doc_type,
        name=name,
        designation=designation,
        version="1.0",
        branch_name=branch_name,
        fields_json=fields_json,
    )
    db.add(doc)
    db.commit()
    db.refresh(doc)

    # ── Сохранить оригинальный файл в хранилище версий ───────────────
    try:
        doc_storage.save_document_version(
            doc_id=doc_id, doc_name=doc.name, doc_type=doc_type,
            version="1.0", docx_bytes=content,
            author=meta_info.get("author", "upload"),
            reason="Первоначальная загрузка файла",
            file_ext=ext,   # .docx | .pdf | .xlsx
        )
    except Exception:
        pass  # Не критично — только файловое хранилище

    # Считаем сколько DSR-полей заполнено
    dsr_fields = [f["id"] for f in (dt.fields_json or [])]
    filled = [f for f in dsr_fields if fields_json.get(f)]

    return {
        "id": doc.id,
        "name": doc.name,
        "designation": doc.designation,
        "doc_type": doc.doc_type,
        "format": ext.lstrip("."),
        "dsr_fields_total": len(dsr_fields),
        "dsr_fields_filled": len(filled),
        "filled_fields": {f: fields_json[f] for f in filled},
        "auto_detected_type": meta_info.get("_auto_detected_type"),
        "pages": meta_info.get("pages"),
        "tables": meta_info.get("tables", 0),
        "table_rows": meta_info.get("table_rows", meta_info.get("total_data_rows", 0)),
        "parse_warning": meta_info.get("parse_warning"),
        "message": (
            f"Файл «{file.filename}» загружен · документ {doc_id} создан · "
            f"заполнено {len(filled)}/{len(dsr_fields)} полей DSR"
        ),
    }


@app.get("/api/documents/{doc_id}", response_model=schemas.DocumentOut, tags=["Документы"])
def get_document(doc_id: str, db: Session = Depends(get_db)):
    doc = db.query(models.DocumentInstance).filter_by(id=doc_id).first()
    if not doc:
        raise HTTPException(404, "Документ не найден")
    return doc


@app.patch("/api/documents/{doc_id}", response_model=schemas.DocumentOut, tags=["Документы"])
def update_document(doc_id: str, data: schemas.DocumentUpdate, db: Session = Depends(get_db)):
    """Частичное обновление документа: наименование, обозначение, версия, статус"""
    doc = db.query(models.DocumentInstance).filter_by(id=doc_id).first()
    if not doc:
        raise HTTPException(404, "Документ не найден")
    if data.name is not None:
        doc.name = data.name
    if data.designation is not None:
        doc.designation = data.designation
    if data.version is not None:
        doc.version = data.version
    if data.status is not None:
        if data.status not in ("active", "pending", "archived"):
            raise HTTPException(400, "Статус должен быть: active | pending | archived")
        doc.status = data.status
    doc.updated_at = datetime.utcnow()
    db.commit()
    db.refresh(doc)
    # Пересохранить DOCX с обновлёнными полями
    _generate_and_store(doc, reason="Редактирование метаданных документа")
    return doc


@app.delete("/api/documents/{doc_id}", tags=["Документы"])
def delete_document(doc_id: str, db: Session = Depends(get_db)):
    """
    Удаление документа со всеми связанными дельтами и уведомлениями.
    ВНИМАНИЕ: необратимая операция.
    """
    doc = db.query(models.DocumentInstance).filter_by(id=doc_id).first()
    if not doc:
        raise HTTPException(404, "Документ не найден")

    # Каскадное удаление: уведомления → дельты → документ
    delta_ids = [d.id for d in db.query(models.Delta).filter_by(doc_id=doc_id).all()]
    if delta_ids:
        db.query(models.Notification).filter(
            models.Notification.delta_id.in_(delta_ids)
        ).delete(synchronize_session=False)
        db.query(models.Delta).filter_by(doc_id=doc_id).delete(synchronize_session=False)
    db.delete(doc)
    db.commit()
    # Удалить физическое хранилище версий
    doc_storage.delete_doc_storage(doc_id)
    return {"status": "deleted", "id": doc_id}


@app.get("/api/documents/{doc_id}/fields", tags=["Документы"])
def get_document_fields(doc_id: str, db: Session = Depends(get_db)):
    """Текущие значения полей документа с метаинформацией из DSR"""
    doc = db.query(models.DocumentInstance).filter_by(id=doc_id).first()
    if not doc:
        raise HTTPException(404, "Документ не найден")
    dt = db.query(models.DocumentType).filter_by(code=doc.doc_type).first()
    schema_fields = dt.fields_json if dt else []
    result = []
    for sf in schema_fields:
        fid = sf["id"]
        deps = [r for r in _matrix
                if r["source_type"] == doc.doc_type and r["source_field"] == fid]
        result.append({
            "id": fid,
            "name": sf["name"],
            "type": sf.get("type", "string"),
            "required": sf.get("required", False),
            "current_value": doc.fields_json.get(fid, ""),
            "dependencies_count": len(deps),
            "dependencies": [{"target": d["target_type"], "target_field": d["target_field"],
                               "dep_type": d["dep_type"]} for d in deps],
        })
    return {"doc_id": doc_id, "doc_type": doc.doc_type, "fields": result}


@app.get("/api/documents/{doc_id}/blame/{field_id}", tags=["Документы"])
def get_blame(doc_id: str, field_id: str, db: Session = Depends(get_db)):
    """История изменений конкретного поля (аналог git blame)"""
    history = blame(db, doc_id, field_id)
    return [schemas.DeltaOut.model_validate(d) for d in history]


@app.get("/api/documents/{doc_id}/checkout", tags=["Документы"])
def get_checkout(doc_id: str,
                 timestamp: datetime = Query(..., description="ISO 8601"),
                 db: Session = Depends(get_db)):
    """Восстановление состояния документа на момент времени"""
    return checkout(db, doc_id, timestamp)


# ===========================================================================
# API: Дельты (Delta Repository)
# ===========================================================================

@app.get("/api/deltas", response_model=List[schemas.DeltaOut], tags=["Дельты"])
def list_deltas(
    doc_id:  Optional[str] = Query(None),
    omega:   Optional[str] = Query(None),
    author:  Optional[str] = Query(None),
    branch:  Optional[str] = Query(None),
    limit:   int = Query(100, le=1000),
    offset:  int = Query(0),
    db: Session = Depends(get_db),
):
    q = db.query(models.Delta)
    if doc_id:  q = q.filter_by(doc_id=doc_id)
    if omega:   q = q.filter_by(omega_type=omega)
    if author:  q = q.filter(models.Delta.author.contains(author))
    if branch:  q = q.filter_by(branch_name=branch)
    return q.order_by(models.Delta.committed_at.desc()).offset(offset).limit(limit).all()


@app.post("/api/deltas", response_model=schemas.DeltaOut, status_code=201, tags=["Дельты"])
def create_delta(data: schemas.DeltaCreate, db: Session = Depends(get_db)):
    """
    Фиксация новой атомарной дельты.
    Автоматически: вычисление SHA-256, обновление документа,
    расчёт каскада E(δ), создание задач-уведомлений.
    """
    # Подставить имя поля из DSR, если не передано
    field_name = data.field_name
    if not field_name:
        doc = db.query(models.DocumentInstance).filter_by(id=data.doc_id).first()
        if doc:
            dt = db.query(models.DocumentType).filter_by(code=doc.doc_type).first()
            if dt:
                for f in dt.fields_json:
                    if f["id"] == data.field_id:
                        field_name = f["name"]
                        break
        field_name = field_name or data.field_id

    try:
        delta = commit_delta(
            db=db,
            doc_id=data.doc_id,
            field_id=data.field_id,
            field_name=field_name,
            v_before=data.v_before,
            v_after=data.v_after,
            omega=data.omega_type,
            author=data.author,
            reason=data.reason,
            iin=data.iin,
            branch=data.branch,
            parent_ids=data.parent_ids,
        )
    except ValueError as e:
        raise HTTPException(409, str(e))

    # Движок зависимостей (с учётом стадии ЖЦ — раздел 4.3)
    try:
        process_delta(db, delta, project_stage=data.project_stage)
    except Exception as exc:
        db.rollback()
        raise HTTPException(500, f"Ошибка движка зависимостей: {exc}")
    db.commit()
    db.refresh(delta)

    # Генерировать и сохранить новую версию DOCX в хранилище
    _save_delta_version(db, delta)

    return delta


@app.get("/api/deltas/{delta_id}", response_model=schemas.DeltaOut, tags=["Дельты"])
def get_delta(delta_id: str, db: Session = Depends(get_db)):
    d = db.query(models.Delta).filter_by(id=delta_id).first()
    if not d:
        raise HTTPException(404, "Дельта не найдена")
    return d


@app.get("/api/deltas/{delta_id}/cascade", tags=["Дельты"])
def get_cascade_for_delta(delta_id: str, db: Session = Depends(get_db)):
    """Полный граф каскада для дельты"""
    return get_full_cascade_tree(delta_id, db)


@app.get("/api/documents/{doc_id}/diff", tags=["Дельты"])
def get_diff(
    doc_id: str,
    t1: Optional[datetime] = Query(None),
    t2: Optional[datetime] = Query(None),
    db: Session = Depends(get_db),
):
    """Все дельты документа за период (аналог git diff)"""
    deltas = diff(db, doc_id, t1, t2)
    return [schemas.DeltaOut.model_validate(d) for d in deltas]


# ===========================================================================
# API: Ветки (Branch)
# ===========================================================================

@app.get("/api/branches", response_model=List[schemas.BranchOut], tags=["Ветки"])
def list_branches(db: Session = Depends(get_db)):
    return db.query(models.Branch).all()


@app.post("/api/branches", response_model=schemas.BranchOut, status_code=201, tags=["Ветки"])
def create_branch_api(data: schemas.BranchCreate, db: Session = Depends(get_db)):
    try:
        br = create_branch(db, data.name, data.description, data.base)
        db.commit()
        db.refresh(br)
        return br
    except ValueError as e:
        raise HTTPException(409, str(e))


@app.post("/api/branches/merge", tags=["Ветки"])
def merge_branches(req: schemas.MergeRequest, db: Session = Depends(get_db)):
    """Слияние двух веток с обнаружением конфликтов"""
    result = merge_branch(db, req.source, req.target, req.resolver)
    db.commit()
    return result


# ===========================================================================
# API: Уведомления (Dependency Tasks)
# ===========================================================================

@app.get("/api/notifications", response_model=List[schemas.NotificationOut], tags=["Уведомления"])
def list_notifications(
    status:    Optional[str] = Query(None),
    dep_type:  Optional[str] = Query(None),
    doc_type:  Optional[str] = Query(None),
    db: Session = Depends(get_db),
):
    q = db.query(models.Notification)
    if status:   q = q.filter_by(status=status)
    if dep_type: q = q.filter_by(dep_type=dep_type)
    if doc_type: q = q.filter_by(target_doc_type=doc_type)
    return q.order_by(models.Notification.created_at.desc()).all()


@app.put("/api/notifications/{notif_id}/resolve", tags=["Уведомления"])
def resolve_notification(notif_id: int, data: schemas.NotificationResolve,
                         db: Session = Depends(get_db)):
    n = db.query(models.Notification).filter_by(id=notif_id).first()
    if not n:
        raise HTTPException(404, "Уведомление не найдено")
    n.status = "resolved"
    n.resolved_at = datetime.utcnow()
    n.notes = data.notes
    # Табл. 4.3: событие NotificationCompleted
    from delta_repo import publish_event
    publish_event(db, "NotificationCompleted", {
        "notif_id": notif_id,
        "completed_by": getattr(data, "completed_by", ""),
        "delta_id": n.delta_id,
        "target_doc_type": n.target_doc_type,
        "target_field": n.target_field,
    })
    db.commit()
    return {"status": "resolved", "id": notif_id}


@app.put("/api/notifications/{notif_id}/skip", tags=["Уведомления"])
def skip_notification(notif_id: int, db: Session = Depends(get_db)):
    n = db.query(models.Notification).filter_by(id=notif_id).first()
    if not n:
        raise HTTPException(404, "Уведомление не найдено")
    n.status = "skipped"
    n.resolved_at = datetime.utcnow()
    db.commit()
    return {"status": "skipped", "id": notif_id}


# ===========================================================================
# API: Конфликты
# ===========================================================================

@app.get("/api/conflicts", tags=["Конфликты"])
def list_conflicts(db: Session = Depends(get_db)):
    return db.query(models.ConflictRecord).filter_by(status="open").all()


@app.post("/api/conflicts/resolve", tags=["Конфликты"])
def resolve_conflict(data: schemas.ConflictResolve, db: Session = Depends(get_db)):
    cf = db.query(models.ConflictRecord).filter_by(id=data.conflict_id).first()
    if not cf:
        raise HTTPException(404, "Конфликт не найден")
    cf.resolved_value = data.resolved_value
    cf.resolver = data.resolver
    cf.resolved_at = datetime.utcnow()
    cf.status = "resolved"
    # Табл. 4.3: событие ConflictResolved
    from delta_repo import publish_event
    publish_event(db, "ConflictResolved", {
        "conflict_id": data.conflict_id,
        "resolved_by": data.resolver,
        "resolved_value": data.resolved_value,
        "doc_id": cf.doc_id,
        "field_id": cf.field_id,
    })
    db.commit()
    return {"status": "resolved"}


# ===========================================================================
# API: Матрица зависимостей
# ===========================================================================

@app.get("/api/matrix", tags=["Матрица зависимостей"])
def get_full_matrix(
    source_type:  Optional[str] = Query(None),
    target_type:  Optional[str] = Query(None),
    dep_type:     Optional[str] = Query(None),
    omega:        Optional[str] = Query(None),
):
    """Полная матрица M[DocType][Field][Ω] с фильтрацией"""
    rules = _matrix
    if source_type: rules = [r for r in rules if r["source_type"] == source_type]
    if target_type: rules = [r for r in rules if r["target_type"] == target_type]
    if dep_type:    rules = [r for r in rules if r["dep_type"] == dep_type]
    if omega:       rules = [r for r in rules if omega in r["omega_types"]]
    return {"rules": rules, "count": len(rules)}


@app.get("/api/matrix/stats", tags=["Матрица зависимостей"])
def matrix_stats():
    return get_matrix_stats()


@app.post("/api/admin/recalculate-cascade", tags=["Администрирование"])
def recalculate_cascade(db: Session = Depends(get_db)):
    """
    Пересчёт каскадных уведомлений для всех дельт.
    Идемпотентен: уже существующие уведомления не дублируются — создаются
    только уведомления для правил матрицы, у которых ещё нет соответствующей записи.
    Применяется после обновления матрицы или добавления демо-данных.
    """
    all_deltas = db.query(models.Delta).all()
    total_new = 0
    deltas_updated = 0

    for delta in all_deltas:
        new_notifs = process_delta(db, delta)
        if new_notifs:
            total_new += len(new_notifs)
            deltas_updated += 1

    db.commit()
    return {
        "message": (
            f"Пересчёт завершён: создано {total_new} новых уведомлений "
            f"для {deltas_updated} дельт. "
            f"Уже существовавшие уведомления не затронуты."
        ),
        "new_notifications": total_new,
        "deltas_updated": deltas_updated,
        "total_deltas": len(all_deltas),
    }


@app.post("/api/matrix/reload", tags=["Матрица зависимостей"])
def reload_matrix_api(db: Session = Depends(get_db)):
    """Перезагрузка матрицы и DSR из файла (без перезапуска). Публикует событие MatrixReloaded."""
    result = reload_matrix()
    from delta_repo import publish_event
    publish_event(db, "MatrixReloaded", {
        "rules_count": result["rules"],
        "mandatory": result["mandatory"],
        "recommended": result["recommended"],
        "dsr_types": result["dsr_types"],
        "timestamp": datetime.utcnow().isoformat(),
    })
    db.commit()
    return result


@app.get("/api/matrix/lifecycle", tags=["Матрица зависимостей"])
def get_lifecycle_map():
    """
    Карта стадий ЖЦ и omega_sensitivity для всех типов документов (раздел 4.3).
    Используется для фильтрации каскада: уведомление не создаётся,
    если целевой документ не обязателен на текущей стадии проекта.
    Стадии: ТП | РД | ИЗГОТ | ЭКСПЛ
    """
    from dependency_engine import _dsr
    return {
        code: {
            "name": info.get("name", ""),
            "std":  info.get("std", ""),
            "lifecycle":        info.get("lifecycle", []),
            "omega_sensitivity": info.get("omega_sensitivity", []),
        }
        for code, info in _dsr.items()
    }


# ===========================================================================
# API: IA Module — интеллектуальный анализ
# ===========================================================================

@app.post("/api/ia/classify", tags=["IA Module"])
def classify_omega_api(req: schemas.ClassifyRequest):
    """Детерминированная классификация типа Ω"""
    result = classify_omega(req.field_id, req.v_before, req.v_after, req.doc_type)
    sig = estimate_significance(result["suggested_omega"], 0)
    return {**result, **sig}


@app.post("/api/ia/significance", tags=["IA Module"])
def significance_api(req: schemas.SignificanceRequest):
    """
    Оценка существенности изменения по формуле §4.6.2:
    S(δ) = Ω_base + min(1.0, cascade_count / 3)
    Ω₁→0, Ω₂→1, Ω₃/Ω₄/Ω₅→2, Ω₆→3, Ω₇→4
    """
    return estimate_significance(req.omega_type, req.cascade_count)


@app.post("/api/ia/analyze", tags=["IA Module"])
def analyze_delta(req: schemas.AnalyzeRequest, db: Session = Depends(get_db)):
    """
    AI-анализ дельты.
    provider=huggingface (по умолчанию): токен в api_key или env HF_TOKEN
    provider=anthropic: токен в api_key или env ANTHROPIC_API_KEY
    """
    d = db.query(models.Delta).filter_by(id=req.delta_id).first()
    if not d:
        # Попробовать по short_sha
        d = db.query(models.Delta).filter_by(short_sha=req.delta_id).first()
    if not d:
        raise HTTPException(404, "Дельта не найдена")

    doc = db.query(models.DocumentInstance).filter_by(id=d.doc_id).first()
    dt  = db.query(models.DocumentType).filter_by(code=d.doc_type).first()

    delta_data = {
        "doc_name": doc.name if doc else d.doc_id,
        "doc_type": d.doc_type,
        "field_id": d.field_id,
        "field_name": d.field_name,
        "v_before": d.v_before,
        "v_after": d.v_after,
        "omega": d.omega_type,
        "author": d.author,
        "iin": d.iin,
        "reason": d.reason,
    }
    dsr_ctx = {"gost": dt.gost if dt else ""}
    cascade_rules = get_cascade(d.doc_type, d.field_id, d.omega_type)

    import asyncio
    try:
        result = asyncio.run(analyze_delta_ai(
            delta_data, dsr_ctx, cascade_rules,
            api_key=req.api_key,
            provider=req.provider,
            hf_model=req.hf_model,
        ))
    except Exception as exc:
        import traceback
        print(f"[IA ERROR] {traceback.format_exc()}")
        raise HTTPException(500, f"Внутренняя ошибка IA-модуля: {exc}")
    return result


@app.post("/api/ia/step", tags=["IA Module"])
def analyze_step(req: schemas.StepDeltaRequest):
    """Анализ дельты STEP ISO 10303-21"""
    return parse_step_delta(req.content_before, req.content_after)


@app.post("/api/ia/xlsx", tags=["IA Module"])
def analyze_xlsx(req: schemas.XlsxDeltaRequest):
    """Анализ дельты XLSX (спецификации, МК, ведомости)"""
    return parse_xlsx_delta(req.rows_before, req.rows_after, req.field_mapping)


@app.post("/api/ia/docx", tags=["IA Module"])
def analyze_docx(req: schemas.DocxDeltaRequest):
    """Анализ дельты DOCX (ПЗ, ТУ, РЭ, ТЗ)"""
    return parse_docx_delta(req.paragraphs_before, req.paragraphs_after)


# ===========================================================================
# API: Снимок и статистика (Dashboard)
# ===========================================================================

@app.get("/api/stats", tags=["Дашборд"])
def get_stats(db: Session = Depends(get_db)):
    total_docs    = db.query(models.DocumentInstance).count()
    total_deltas  = db.query(models.Delta).count()
    pending_tasks = db.query(models.Notification).filter_by(status="pending").count()
    mandatory_pending = db.query(models.Notification).filter_by(
        status="pending", dep_type="1").count()
    resolved      = db.query(models.Notification).filter_by(status="resolved").count()
    conflicts     = db.query(models.ConflictRecord).filter_by(status="open").count()
    branches      = db.query(models.Branch).filter_by(is_merged=False).count()

    omega_dist = {}
    for omega in ["Ω₁","Ω₂","Ω₃","Ω₄","Ω₅","Ω₆","Ω₇"]:
        omega_dist[omega] = db.query(models.Delta).filter_by(omega_type=omega).count()

    recent_deltas = (
        db.query(models.Delta)
        .order_by(models.Delta.committed_at.desc())
        .limit(5).all()
    )

    pending_notifications = (
        db.query(models.Notification)
        .filter_by(status="pending")
        .order_by(models.Notification.created_at.desc())
        .limit(5).all()
    )

    matrix_s = get_matrix_stats()

    return {
        "documents":         total_docs,
        "deltas":            total_deltas,
        "pending_tasks":     pending_tasks,
        "mandatory_pending": mandatory_pending,
        "resolved_tasks":    resolved,
        "open_conflicts":    conflicts,
        "active_branches":   branches,
        "omega_distribution": omega_dist,
        "matrix_rules":      matrix_s["total_rules"],
        "recent_deltas":     [schemas.DeltaOut.model_validate(d) for d in recent_deltas],
        "urgent_notifications": [schemas.NotificationOut.model_validate(n)
                                 for n in pending_notifications],
    }


@app.get("/api/snapshot", tags=["Дашборд"])
def get_snapshot(
    timestamp: Optional[datetime] = Query(None),
    db: Session = Depends(get_db),
):
    """Полный снимок состояния репозитория"""
    return snapshot(db, timestamp)


@app.get("/api/audit", tags=["Дашборд"])
def get_audit(
    limit:      int            = Query(50, le=500),
    event_type: Optional[str]  = Query(None, description="Фильтр по типу события"),
    delta_id:   Optional[str]  = Query(None, description="Фильтр по ID дельты"),
    from_dt:    Optional[datetime] = Query(None, alias="from", description="ISO 8601 начало периода"),
    to_dt:      Optional[datetime] = Query(None, alias="to",   description="ISO 8601 конец периода"),
    db: Session = Depends(get_db),
):
    """Журнал событий шины (Event Bus) с фильтрацией (раздел 4.5)"""
    q = db.query(models.AuditEvent)
    if event_type:
        q = q.filter(models.AuditEvent.event_type.contains(event_type))
    if from_dt:
        q = q.filter(models.AuditEvent.occurred_at >= from_dt)
    if to_dt:
        q = q.filter(models.AuditEvent.occurred_at <= to_dt)
    events = q.order_by(models.AuditEvent.occurred_at.desc()).limit(limit * 5 if delta_id else limit).all()
    result = []
    for e in events:
        # Фильтрация delta_id на уровне Python (JSON path не поддерживается в SQLite < 3.38)
        if delta_id and e.payload.get("delta_id") != delta_id:
            continue
        result.append({"id": e.id, "event_type": e.event_type,
                        "payload": e.payload, "occurred_at": e.occurred_at})
        if len(result) >= limit:
            break
    return result


# ===========================================================================
# API: Версионное хранилище файлов
# ===========================================================================

@app.get("/api/documents/{doc_id}/versions", tags=["Версии файлов"])
def list_versions(doc_id: str, db: Session = Depends(get_db)):
    """
    Список версий DOCX-файлов документа в хранилище (аналог git log --oneline).
    Структура: prj_docs/{doc_id}/v{version}.docx
    """
    doc = db.query(models.DocumentInstance).filter_by(id=doc_id).first()
    if not doc:
        raise HTTPException(404, "Документ не найден")
    stats = doc_storage.get_storage_stats(doc_id)
    versions = doc_storage.list_document_versions(doc_id)
    return {
        "doc_id": doc_id,
        "head": stats.get("head"),
        "version_count": stats.get("version_count", 0),
        "storage_path": stats.get("storage_path"),
        "versions": versions,
    }


@app.get("/api/documents/{doc_id}/file", tags=["Версии файлов"])
def download_version(
    doc_id:  str,
    version: Optional[str] = Query(None, description="Версия файла. По умолчанию HEAD"),
    db: Session = Depends(get_db),
):
    """
    Скачать DOCX-файл документа (текущую или указанную версию).
    Если файл не сохранён — генерирует его на лету из полей БД.
    """
    doc = db.query(models.DocumentInstance).filter_by(id=doc_id).first()
    if not doc:
        raise HTTPException(404, "Документ не найден")

    # Пробуем взять из хранилища
    data = doc_storage.get_version_bytes(doc_id, version)
    file_ext = doc_storage.get_version_format(doc_id, version)

    if data is None:
        # Генерируем DOCX на лету
        try:
            from generate_docs import generate_docx
            data = generate_docx(doc)
            file_ext = ".docx"
        except Exception as e:
            raise HTTPException(500, f"Не удалось сгенерировать файл: {e}")

    ver = version or doc_storage.get_head_version(doc_id) or doc.version
    filename = f"{doc_id}_v{ver}{file_ext}"

    _MEDIA_TYPES = {
        ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ".pdf":  "application/pdf",
        ".xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    }
    media_type = _MEDIA_TYPES.get(file_ext,
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document")

    return Response(
        content=data,
        media_type=media_type,
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@app.post("/api/documents/{doc_id}/generate", tags=["Версии файлов"])
def generate_and_save(
    doc_id: str,
    author: str = Query("system"),
    db: Session = Depends(get_db),
):
    """
    Генерирует DOCX из текущих полей документа и сохраняет в хранилище.
    Возвращает путь к файлу.
    """
    doc = db.query(models.DocumentInstance).filter_by(id=doc_id).first()
    if not doc:
        raise HTTPException(404, "Документ не найден")
    try:
        from generate_docs import generate_docx
        data = generate_docx(doc)
    except Exception as e:
        raise HTTPException(500, f"Ошибка генерации: {e}")

    path = doc_storage.save_document_version(
        doc_id=doc_id, doc_name=doc.name, doc_type=doc.doc_type,
        version=doc.version, docx_bytes=data,
        author=author, reason="Ручная генерация из UI",
        committed_at=datetime.utcnow(),
    )
    return {
        "status": "saved",
        "doc_id": doc_id,
        "version": doc.version,
        "path": path,
        "size_bytes": len(data),
    }


# ---------------------------------------------------------------------------
# Вспомогательная функция: сохранить версию после дельты
# ---------------------------------------------------------------------------

def _generate_and_store(doc, reason: str = "", author: str = "system",
                        delta_id: str = None, committed_at=None):
    """
    Генерирует DOCX для документа и сохраняет его в хранилище версий.
    Не бросает исключений — ошибка файлового хранилища не должна ломать API.
    """
    try:
        from generate_docs import generate_docx
        data = generate_docx(doc)
        doc_storage.save_document_version(
            doc_id=doc.id,
            doc_name=doc.name,
            doc_type=doc.doc_type,
            version=doc.version,
            docx_bytes=data,
            delta_id=delta_id,
            author=author,
            reason=reason,
            committed_at=committed_at or datetime.utcnow(),
        )
    except Exception as e:
        print(f"[doc_storage] {doc.id}: {e}")


def _save_delta_version(db: Session, delta):
    """После фиксации дельты пересохраняем DOCX с обновлёнными полями."""
    doc = db.query(models.DocumentInstance).filter_by(id=delta.doc_id).first()
    if doc is None:
        return
    _generate_and_store(
        doc,
        reason=delta.reason or f"Дельта {delta.short_sha}",
        author=delta.author,
        delta_id=delta.id,
        committed_at=delta.committed_at,
    )
